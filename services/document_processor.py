from __future__ import annotations

import io
import os
from dataclasses import dataclass, field
from typing import Any


@dataclass(slots=True)
class ContentUnderstandingSettings:
    endpoint: str = ""
    key: str = ""
    analyzer_id: str = "prebuilt-documentSearch"
    api_version: str = "2025-11-01"

    @classmethod
    def from_env(cls) -> "ContentUnderstandingSettings":
        return cls(
            endpoint=os.getenv("CONTENTUNDERSTANDING_ENDPOINT", "")
            or os.getenv("AZURE_CONTENT_UNDERSTANDING_ENDPOINT", ""),
            key=os.getenv("CONTENTUNDERSTANDING_KEY", "")
            or os.getenv("CONTENT_UNDERSTANDING_KEY", ""),
            analyzer_id=os.getenv("CONTENTUNDERSTANDING_ANALYZER_ID", "prebuilt-documentSearch"),
            api_version=os.getenv("CONTENTUNDERSTANDING_API_VERSION", "2025-11-01"),
        )

    @property
    def is_configured(self) -> bool:
        return bool(self.endpoint.strip())


@dataclass(slots=True)
class DocumentProcessingResult:
    text: str
    source: str
    file_name: str
    raw_result: dict[str, Any] | None = None
    warnings: list[str] = field(default_factory=list)


class DocumentProcessor:
    """Extracts resume text with Azure Content Understanding, with local fallback parsers."""

    def __init__(self, settings: ContentUnderstandingSettings | None = None) -> None:
        self.settings = settings or ContentUnderstandingSettings.from_env()

    def process_uploaded_file(self, uploaded_file: Any) -> DocumentProcessingResult:
        file_name = getattr(uploaded_file, "name", "resume")
        file_bytes = uploaded_file.getvalue()
        warnings: list[str] = []

        if self.settings.is_configured:
            try:
                text, raw_result = self._extract_with_content_understanding(file_bytes)
                if text.strip():
                    return DocumentProcessingResult(
                        text=text,
                        source="Azure Content Understanding",
                        file_name=file_name,
                        raw_result=raw_result,
                        warnings=warnings,
                    )
                warnings.append("Azure Content Understanding returned no text; local extraction was used.")
            except Exception as exc:  # noqa: BLE001 - surfaced to UI as a non-fatal extraction warning.
                warnings.append(f"Azure Content Understanding failed: {exc}")
        else:
            warnings.append("Content Understanding endpoint is not configured; local extraction was used.")

        local_text = self._extract_locally(file_name=file_name, file_bytes=file_bytes)
        return DocumentProcessingResult(
            text=local_text,
            source="Local parser",
            file_name=file_name,
            raw_result=None,
            warnings=warnings,
        )

    def _extract_with_content_understanding(self, file_bytes: bytes) -> tuple[str, dict[str, Any] | None]:
        from azure.ai.contentunderstanding import ContentUnderstandingClient, to_llm_input
        from azure.core.credentials import AzureKeyCredential
        from azure.identity import DefaultAzureCredential

        credential = AzureKeyCredential(self.settings.key) if self.settings.key else DefaultAzureCredential()
        client = ContentUnderstandingClient(
            endpoint=self.settings.endpoint,
            credential=credential,
            api_version=self.settings.api_version,
        )
        poller = client.begin_analyze_binary(
            analyzer_id=self.settings.analyzer_id,
            binary_input=file_bytes,
        )
        result = poller.result()

        raw_result = result.as_dict() if hasattr(result, "as_dict") else None
        try:
            text = to_llm_input(result)
        except Exception:  # noqa: BLE001 - SDK helper may not exist in older package builds.
            text = self._extract_text_from_content_result(result)
        return text, raw_result

    def _extract_text_from_content_result(self, result: Any) -> str:
        chunks: list[str] = []
        for content in getattr(result, "contents", []) or []:
            markdown = getattr(content, "markdown", "")
            if markdown:
                chunks.append(markdown)
            text = getattr(content, "text", "")
            if text and text not in chunks:
                chunks.append(text)
        if chunks:
            return "\n\n".join(chunks)
        if hasattr(result, "as_dict"):
            return str(result.as_dict())
        return str(result)

    def _extract_locally(self, file_name: str, file_bytes: bytes) -> str:
        suffix = file_name.lower().rsplit(".", maxsplit=1)[-1] if "." in file_name else ""
        if suffix == "pdf":
            return self._extract_pdf_text(file_bytes)
        if suffix == "docx":
            return self._extract_docx_text(file_bytes)
        raise ValueError("Unsupported resume format. Upload a PDF or DOCX file.")

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        import pdfplumber

        chunks: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text.strip())
        return "\n\n".join(chunks)

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells: list[str] = []
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    table_cells.append(" | ".join(values))
        return "\n".join([*paragraphs, *table_cells])
